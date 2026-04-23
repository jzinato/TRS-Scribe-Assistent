import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

import streamlit as st

try:
    from PIL import Image
except Exception:
    Image = None

# Optional imports. The app works without them.
try:
    import docx
except Exception:
    docx = None

# ==============================
# Page config
# ==============================
st.set_page_config(
    page_title="TRS Scribe Assist",
    page_icon="🩺",
    layout="wide",
)

# ==============================
# Constants and dictionaries
# ==============================
EXAM_SYNONYMS = {
    "hemoglobina": ["hb", "hgb", "hemoglobina"],
    "hematocrito": ["ht", "hct", "hematocrito", "hematócrito"],
    "ferritina": ["ferritina"],
    "ist": ["ist", "tsat", "saturacao de transferrina", "saturação de transferrina"],
    "ferro_serico": ["ferro", "ferro serico", "ferro sérico"],
    "calcio": ["calcio", "cálcio", "calcio total", "cálcio total"],
    "fosforo": ["fosforo", "fósforo", "p"],
    "pth": ["pth", "paratormonio", "paratormônio"],
    "fosfatase_alcalina": ["fa", "fosfatase alcalina"],
    "vitamina_d": ["vitamina d", "25-oh vitamina d", "25 oh d"],
    "potassio": ["potassio", "potássio", "k"],
    "sodio": ["sodio", "sódio", "na"],
    "bicarbonato": ["bicarbonato", "hco3"],
    "ureia": ["ureia", "urea"],
    "creatinina": ["creatinina"],
    "albumina": ["albumina"],
    "pcr": ["pcr", "proteina c reativa", "proteína c reativa"],
    "hbsag": ["hbsag"],
    "anti_hbs": ["anti hbs", "anti-hbs"],
    "anti_hcv": ["anti hcv", "anti-hcv"],
    "hiv": ["hiv", "anti hiv", "anti-hiv"],
}

IMPORTANT_EXAMS = [
    "hemoglobina", "ferritina", "ist", "calcio", "fosforo", "pth",
    "potassio", "bicarbonato", "albumina", "hbsag", "anti_hbs", "anti_hcv"
]


# ==============================
# Data models
# ==============================
@dataclass
class ExamItem:
    nome_padronizado: str
    nome_original: str
    valor: str
    unidade: str
    referencia: str
    data_coleta: str
    incerto: bool
    trecho_fonte: str


@dataclass
class MedicationItem:
    nome: str
    dose: str
    via: str
    frequencia: str
    fonte: str


@dataclass
class ClinicalCase:
    paciente: Dict[str, Any]
    contexto: Dict[str, Any]
    exames: List[Dict[str, Any]]
    medicacoes_em_uso: List[Dict[str, Any]]
    prescricao_dialitica: Dict[str, Any]
    evolucao_previa: Dict[str, Any]
    intercorrencias: List[Dict[str, Any]]
    problemas_ativos: List[Dict[str, Any]]
    pendencias: List[Dict[str, Any]]
    rastreabilidade: Dict[str, Any]


# ==============================
# Utility functions
# ==============================
def normalize_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "á": "a", "à": "a", "â": "a", "ã": "a",
        "é": "e", "ê": "e",
        "í": "i",
        "ó": "o", "ô": "o", "õ": "o",
        "ú": "u",
        "ç": "c",
    }
    out = text.lower()
    for old, new in replacements.items():
        out = out.replace(old, new)
    return out


def standardize_exam_name(raw_name: str) -> str:
    name_norm = normalize_text(raw_name)
    for standard, synonyms in EXAM_SYNONYMS.items():
        for syn in synonyms:
            if syn in name_norm or name_norm == syn:
                return standard
    return name_norm.replace(" ", "_")


def parse_date(value: str) -> str:
    if not value:
        return ""
    for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d"]:
        try:
            return datetime.strptime(value.strip(), fmt).date().isoformat()
        except Exception:
            continue
    return value.strip()


def basic_extract_exams(text: str, default_date: str = "") -> List[ExamItem]:
    """
    Simple regex-based extractor for pasted text.
    Deliberately conservative. It avoids inventing data and only captures patterns it sees.
    Accepted patterns include lines such as:
      Hemoglobina: 9,8 g/dL
      Fósforo 6.1 mg/dL
      PTH = 780 pg/mL
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    exams: List[ExamItem] = []
    pattern = re.compile(
        r"^(?P<name>[A-Za-zÀ-ÿ0-9\-\s\/%]+?)\s*(?:\:|\=)?\s*"
        r"(?P<value>[<>]?\d+[\.,]?\d*)\s*"
        r"(?P<unit>[A-Za-zµ/%²0-9\-\^\.]+)?\s*$"
    )

    for line in lines:
        m = pattern.match(line)
        if not m:
            continue
        raw_name = m.group("name").strip(" -")
        value = m.group("value") or ""
        unit = (m.group("unit") or "").strip()
        std_name = standardize_exam_name(raw_name)

        exams.append(
            ExamItem(
                nome_padronizado=std_name,
                nome_original=raw_name,
                valor=value.replace(",", "."),
                unidade=unit,
                referencia="",
                data_coleta=default_date,
                incerto=False,
                trecho_fonte=line,
            )
        )
    return exams


def parse_medications(text: str) -> List[MedicationItem]:
    meds: List[MedicationItem] = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for line in lines:
        meds.append(MedicationItem(
            nome=line,
            dose="",
            via="",
            frequencia="",
            fonte="texto_manual"
        ))
    return meds


def exams_to_index(exams: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    indexed: Dict[str, Dict[str, Any]] = {}
    for exam in exams:
        key = exam.get("nome_padronizado", "")
        if key:
            indexed[key] = exam
    return indexed


def try_float(value: str) -> Optional[float]:
    try:
        return float(str(value).replace(",", "."))
    except Exception:
        return None


def reconciler(current_exams: List[Dict[str, Any]], previous_exams: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    current = exams_to_index(current_exams)
    previous = exams_to_index(previous_exams)
    findings: List[Dict[str, Any]] = []

    tracked = sorted(set(current.keys()) | set(previous.keys()))
    for key in tracked:
        curr = current.get(key)
        prev = previous.get(key)
        if curr and prev:
            curr_val = try_float(curr.get("valor", ""))
            prev_val = try_float(prev.get("valor", ""))
            trend = "sem tendencia definida"
            if curr_val is not None and prev_val is not None:
                if curr_val > prev_val:
                    trend = "aumentou"
                elif curr_val < prev_val:
                    trend = "reduziu"
                else:
                    trend = "estavel"
            findings.append({
                "dominio": key,
                "descricao": f"{key}: atual={curr.get('valor','')} {curr.get('unidade','')} | previo={prev.get('valor','')} {prev.get('unidade','')} | tendencia={trend}",
                "status": "comparado"
            })
        elif curr and not prev:
            findings.append({
                "dominio": key,
                "descricao": f"{key}: presente no material atual sem comparativo previo estruturado.",
                "status": "novo_no_contexto"
            })
        elif prev and not curr:
            findings.append({
                "dominio": key,
                "descricao": f"{key}: presente apenas no material previo; nao localizado no material atual.",
                "status": "ausente_no_atual"
            })
    return findings


def guideline_engine(case_data: ClinicalCase) -> List[Dict[str, Any]]:
    exams = exams_to_index(case_data.exames)
    suggestions: List[Dict[str, Any]] = []

    hb = try_float(exams.get("hemoglobina", {}).get("valor", ""))
    ferritin = try_float(exams.get("ferritina", {}).get("valor", ""))
    tsat = try_float(exams.get("ist", {}).get("valor", ""))
    phosphorus = try_float(exams.get("fosforo", {}).get("valor", ""))
    pth = try_float(exams.get("pth", {}).get("valor", ""))
    potassium = try_float(exams.get("potassio", {}).get("valor", ""))

    if hb is not None:
        lacunas = []
        if ferritin is None:
            lacunas.append("ferritina ausente")
        if tsat is None:
            lacunas.append("IST/TSAT ausente")
        suggestions.append({
            "problema": "anemia",
            "achado": f"Hemoglobina atual: {hb}",
            "base_documental": "Base guideline a ser conectada via RAG; placeholder para KDIGO anemia e protocolo local.",
            "lacunas": lacunas,
            "sugestao_prudente": "Revisar estrategia de anemia com correlacao longitudinal e protocolo local. Nao fechar prescricao automaticamente."
        })

    if phosphorus is not None or pth is not None:
        lacunas = []
        if exams.get("calcio") is None:
            lacunas.append("calcio ausente")
        if exams.get("fosfatase_alcalina") is None:
            lacunas.append("fosfatase alcalina ausente")
        suggestions.append({
            "problema": "disturbio_mineral_osseo",
            "achado": f"Fosforo={phosphorus if phosphorus is not None else 'NA'} | PTH={pth if pth is not None else 'NA'}",
            "base_documental": "Base guideline a ser conectada via RAG; placeholder para CKD-MBD e protocolo local.",
            "lacunas": lacunas,
            "sugestao_prudente": "Reavaliar tendencia, adesao e esquema atual antes de qualquer ajuste terapeutico."
        })

    if potassium is not None:
        severity = "valor critico" if potassium >= 6.0 else "sem criticidade automatica"
        suggestions.append({
            "problema": "potassio",
            "achado": f"Potassio atual: {potassium} ({severity})",
            "base_documental": "Placeholder de seguranca; exige revisao medica imediata se criticidade clinica.",
            "lacunas": [],
            "sugestao_prudente": "Checar contexto clinico, hemolise, ECG e necessidade de acao imediata conforme protocolo local."
        })

    return suggestions


def build_short_evolution(case_data: ClinicalCase, guideline_notes: List[Dict[str, Any]]) -> str:
    modality = case_data.contexto.get("modalidade_trs", "TRS nao informada")
    summary_parts = []
    problems = [g["problema"] for g in guideline_notes]

    if "anemia" in problems:
        summary_parts.append("achados relacionados a anemia em seguimento")
    if "disturbio_mineral_osseo" in problems:
        summary_parts.append("alteracoes de metabolismo mineral/osseo")
    if "potassio" in problems:
        summary_parts.append("necessidade de correlacao clinica do potassio")

    if not summary_parts:
        summary_parts.append("sem achados automaticamente categorizados pelo modelo inicial")

    return (
        f"Paciente em {modality}. Material atual mostra {', '.join(summary_parts)}. "
        f"Minuta automatizada para revisao medica, sem fechamento autonomo de prescricao."
    )


def build_detailed_evolution(case_data: ClinicalCase, guideline_notes: List[Dict[str, Any]], reconciliation: List[Dict[str, Any]]) -> str:
    lines = []
    modality = case_data.contexto.get("modalidade_trs", "nao informada")
    lines.append(f"Paciente em {modality}.")

    if guideline_notes:
        lines.append("Achados estruturados:")
        for item in guideline_notes:
            lines.append(f"- {item['problema']}: {item['achado']}")

    if reconciliation:
        lines.append("Comparacao longitudinal:")
        for item in reconciliation[:8]:
            lines.append(f"- {item['descricao']}")

    pend = case_data.pendencias
    if pend:
        lines.append("Pendencias:")
        for item in pend:
            lines.append(f"- {item.get('descricao','')}")

    lines.append("Texto gerado para revisao humana final.")
    return "\n".join(lines)


def safety_auditor(case_data: ClinicalCase, detailed_note: str) -> Dict[str, Any]:
    exams = exams_to_index(case_data.exames)
    issues: List[str] = []

    if not case_data.contexto.get("modalidade_trs"):
        issues.append("Modalidade de TRS ausente.")

    for name in ["hemoglobina", "potassio", "fosforo", "calcio", "pth"]:
        exam = exams.get(name)
        if exam and not exam.get("unidade"):
            issues.append(f"Exame critico sem unidade: {name}.")

    potassium = try_float(exams.get("potassio", {}).get("valor", ""))
    if potassium is not None and potassium >= 6.0 and "potassio" not in normalize_text(detailed_note):
        issues.append("Possivel hiperpotassemia sem destaque adequado na narrativa.")

    if issues:
        status = "Bloqueado" if len(issues) >= 2 else "Aprovado com ressalvas"
    else:
        status = "Aprovado"

    return {
        "status": status,
        "ressalvas": issues,
    }


def default_case() -> ClinicalCase:
    return ClinicalCase(
        paciente={"nome": "", "registro": "", "data_nascimento": "", "sexo": ""},
        contexto={
            "modalidade_trs": "HD",
            "local_atendimento": "",
            "data_referencia": datetime.now().date().isoformat(),
            "fonte_entrada": [],
        },
        exames=[],
        medicacoes_em_uso=[],
        prescricao_dialitica={
            "modalidade": "",
            "frequencia": "",
            "duracao_sessao": "",
            "banho": "",
            "anticoagulacao": "",
            "uf_programada": "",
            "peso_seco": "",
            "acesso": "",
        },
        evolucao_previa={"data": "", "texto": ""},
        intercorrencias=[],
        problemas_ativos=[],
        pendencias=[],
        rastreabilidade={
            "documentos_entrada": [],
            "guidelines_consultadas": [],
            "versao_modelo": "mvp-streamlit-v0.1",
            "data_processamento": datetime.now().isoformat(),
        },
    )


def add_basic_pendencies(case_data: ClinicalCase) -> None:
    exams = exams_to_index(case_data.exames)
    pending = []
    for key in IMPORTANT_EXAMS:
        if key not in exams:
            pending.append({"tipo": "exame_ausente", "descricao": f"Exame nao localizado: {key}"})
    case_data.pendencias = pending


def export_docx(short_note: str, detailed_note: str, audit: Dict[str, Any]) -> Optional[bytes]:
    if docx is None:
        return None

    document = docx.Document()
    document.add_heading("TRS Scribe Assist - Relatorio", level=1)
    document.add_paragraph(f"Status da auditoria: {audit.get('status', '')}")

    document.add_heading("Evolucao curta", level=2)
    document.add_paragraph(short_note)

    document.add_heading("Evolucao detalhada", level=2)
    document.add_paragraph(detailed_note)

    if audit.get("ressalvas"):
        document.add_heading("Ressalvas", level=2)
        for item in audit["ressalvas"]:
            document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ==============================
# Session state
# ==============================
if "case_data" not in st.session_state:
    st.session_state.case_data = default_case()
if "previous_exams" not in st.session_state:
    st.session_state.previous_exams = []
if "results" not in st.session_state:
    st.session_state.results = {}


# ==============================
# UI
# ==============================
st.title("TRS Scribe Assist")
st.caption("MVP inicial para organizacao de exames, reconciliacao, minuta de evolucao e auditoria basica.")

with st.sidebar:
    st.header("Configuracao do caso")
    paciente_nome = st.text_input("Nome do paciente", value=st.session_state.case_data.paciente.get("nome", ""))
    registro = st.text_input("Registro", value=st.session_state.case_data.paciente.get("registro", ""))
    modalidade = st.selectbox("Modalidade TRS", ["HD", "HDF", "DP", "Nao informada"], index=0)
    local = st.text_input("Local de atendimento", value=st.session_state.case_data.contexto.get("local_atendimento", ""))
    data_ref = st.date_input("Data de referencia", value=datetime.now().date())

    st.session_state.case_data.paciente["nome"] = paciente_nome
    st.session_state.case_data.paciente["registro"] = registro
    st.session_state.case_data.contexto["modalidade_trs"] = modalidade
    st.session_state.case_data.contexto["local_atendimento"] = local
    st.session_state.case_data.contexto["data_referencia"] = data_ref.isoformat()


tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Entrada",
    "Extracao",
    "Analise",
    "Saida",
    "Auditoria",
])

with tab1:
    st.subheader("Entrada")
    col1, col2 = st.columns(2)

    with col1:
        uploaded_pdf = st.file_uploader("Upload PDF (placeholder de ingestao)", type=["pdf"], key="pdf")
        uploaded_image = st.file_uploader("Upload imagem", type=["png", "jpg", "jpeg", "webp"], key="img")

        if uploaded_image is not None and Image is not None:
            try:
                img = Image.open(uploaded_image)
                st.image(img, caption="Imagem carregada", use_container_width=True)
            except Exception:
                st.warning("Nao foi possivel renderizar a imagem.")

        if uploaded_pdf is not None:
            st.info("Leitura real de PDF ainda deve ser conectada ao parser/OCR do projeto. Neste MVP o PDF fica apenas registrado.")
            st.session_state.case_data.rastreabilidade["documentos_entrada"].append(uploaded_pdf.name)
            if "pdf" not in st.session_state.case_data.contexto["fonte_entrada"]:
                st.session_state.case_data.contexto["fonte_entrada"].append("pdf")

        if uploaded_image is not None:
            st.session_state.case_data.rastreabilidade["documentos_entrada"].append(uploaded_image.name)
            if "imagem" not in st.session_state.case_data.contexto["fonte_entrada"]:
                st.session_state.case_data.contexto["fonte_entrada"].append("imagem")

    with col2:
        raw_exam_text = st.text_area(
            "Texto de exames atual",
            height=220,
            placeholder="Exemplo:\nHemoglobina: 9,8 g/dL\nFerritina: 180 ng/mL\nFosforo: 6,1 mg/dL",
        )
        previous_exam_text = st.text_area(
            "Texto de exames previos (opcional)",
            height=140,
            placeholder="Exemplo:\nHemoglobina: 10,4 g/dL\nFerritina: 220 ng/mL",
        )

    st.subheader("Contexto clinico adicional")
    meds_text = st.text_area("Medicacoes em uso", height=120, placeholder="Uma por linha")
    prev_note = st.text_area("Evolucao anterior", height=140)
    program_text = st.text_area("Programacao/prescricao previa", height=120)

    if st.button("Processar caso", type="primary"):
        case_data = st.session_state.case_data

        current_exams = [asdict(e) for e in basic_extract_exams(raw_exam_text, case_data.contexto["data_referencia"])]
        previous_exams = [asdict(e) for e in basic_extract_exams(previous_exam_text, "")]
        meds = [asdict(m) for m in parse_medications(meds_text)]

        case_data.exames = current_exams
        case_data.medicacoes_em_uso = meds
        case_data.evolucao_previa = {
            "data": "",
            "texto": prev_note,
        }
        case_data.prescricao_dialitica["modalidade"] = case_data.contexto.get("modalidade_trs", "")
        case_data.prescricao_dialitica["frequencia"] = program_text
        if raw_exam_text and "texto" not in case_data.contexto["fonte_entrada"]:
            case_data.contexto["fonte_entrada"].append("texto")

        add_basic_pendencies(case_data)
        reconciliation = reconciler(case_data.exames, previous_exams)
        notes = guideline_engine(case_data)
        short_note = build_short_evolution(case_data, notes)
        detailed_note = build_detailed_evolution(case_data, notes, reconciliation)
        audit = safety_auditor(case_data, detailed_note)

        st.session_state.case_data = case_data
        st.session_state.previous_exams = previous_exams
        st.session_state.results = {
            "reconciliation": reconciliation,
            "guideline_notes": notes,
            "short_note": short_note,
            "detailed_note": detailed_note,
            "audit": audit,
        }
        st.success("Caso processado.")

with tab2:
    st.subheader("Extracao estruturada")
    case_data = st.session_state.case_data
    if case_data.exames:
        st.write("### Exames atuais")
        st.dataframe(case_data.exames, use_container_width=True)
    else:
        st.info("Nenhum exame estruturado ainda.")

    if st.session_state.previous_exams:
        st.write("### Exames previos")
        st.dataframe(st.session_state.previous_exams, use_container_width=True)

    if case_data.medicacoes_em_uso:
        st.write("### Medicacoes")
        st.dataframe(case_data.medicacoes_em_uso, use_container_width=True)

    st.write("### JSON do caso")
    st.code(json.dumps(asdict(case_data), ensure_ascii=False, indent=2), language="json")

with tab3:
    st.subheader("Analise")
    results = st.session_state.results
    if results.get("reconciliation"):
        st.write("### Reconciliação longitudinal")
        st.dataframe(results["reconciliation"], use_container_width=True)
    else:
        st.info("Analise ainda nao executada.")

    if results.get("guideline_notes"):
        st.write("### Notas do motor de guideline")
        for item in results["guideline_notes"]:
            with st.expander(item["problema"], expanded=True):
                st.write(f"**Achado:** {item['achado']}")
                st.write(f"**Base documental:** {item['base_documental']}")
                st.write(f"**Lacunas:** {', '.join(item['lacunas']) if item['lacunas'] else 'Nenhuma destacada'}")
                st.write(f"**Sugestao prudente:** {item['sugestao_prudente']}")
    else:
        st.info("Sem notas de guideline ainda.")

with tab4:
    st.subheader("Saida")
    results = st.session_state.results
    short_note = results.get("short_note", "")
    detailed_note = results.get("detailed_note", "")

    st.write("### Evolucao curta")
    st.text_area("", short_note, height=100, key="short_out")

    st.write("### Evolucao detalhada")
    st.text_area(" ", detailed_note, height=240, key="long_out")

    if short_note or detailed_note:
        export_json = json.dumps({
            "short_note": short_note,
            "detailed_note": detailed_note,
            "case_data": asdict(st.session_state.case_data),
            "results": st.session_state.results,
        }, ensure_ascii=False, indent=2)

        st.download_button(
            label="Baixar JSON do caso",
            data=export_json.encode("utf-8"),
            file_name="trs_scribe_case.json",
            mime="application/json",
        )

        docx_bytes = export_docx(short_note, detailed_note, results.get("audit", {}))
        if docx_bytes is not None:
            st.download_button(
                label="Baixar DOCX",
                data=docx_bytes,
                file_name="trs_scribe_relatorio.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.caption("python-docx nao encontrado. O botao de DOCX aparece quando a dependencia estiver instalada.")
    else:
        st.info("Processe o caso primeiro.")

with tab5:
    st.subheader("Auditoria")
    audit = st.session_state.results.get("audit", {})
    if audit:
        status = audit.get("status", "")
        if status == "Aprovado":
            st.success(status)
        elif status == "Aprovado com ressalvas":
            st.warning(status)
        else:
            st.error(status)

        ressalvas = audit.get("ressalvas", [])
        if ressalvas:
            st.write("### Ressalvas")
            for item in ressalvas:
                st.write(f"- {item}")
    else:
        st.info("Auditoria ainda nao executada.")

st.divider()
st.caption(
    "Este MVP e um esqueleto operacional. Ele nao substitui revisao medica e nao deve gerar prescricao autonoma. "
    "Os pontos marcados como placeholder sao os locais onde voce conectara OCR, parser de PDF e base RAG de guidelines."
)
