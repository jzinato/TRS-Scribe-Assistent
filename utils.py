import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import docx
except Exception:
    docx = None

EXAM_SYNONYMS = {
    "hemoglobina": ["hb", "hgb", "hemoglobina"],
    "hematocrito": ["ht", "hct", "hematocrito", "hematócrito"],
    "ferritina": ["ferritina"],
    "ist": ["ist", "tsat", "saturacao de transferrina", "saturação de transferrina"],
    "ferro_serico": ["ferro", "ferro serico", "ferro sérico"],
    "calcio": ["calcio", "cálcio", "calcio total", "cálcio total"],
    "fosforo": ["fosforo", "fósforo", "p"],
    "pth": ["pth", "paratormonio", "paratormônio"],
    "fosfatase_alcalina": ["fa", "fosfatase alcalina"],
    "vitamina_d": ["vitamina d", "25-oh vitamina d", "25 oh d"],
    "potassio": ["potassio", "potássio", "k"],
    "sodio": ["sodio", "sódio", "na"],
    "bicarbonato": ["bicarbonato", "hco3"],
    "ureia": ["ureia", "urea"],
    "creatinina": ["creatinina"],
    "albumina": ["albumina"],
    "pcr": ["pcr", "proteina c reativa", "proteína c reativa"],
    "hbsag": ["hbsag"],
    "anti_hbs": ["anti hbs", "anti-hbs"],
    "anti_hcv": ["anti hcv", "anti-hcv"],
    "hiv": ["hiv", "anti hiv", "anti-hiv"],
}

IMPORTANT_EXAMS = [
    "hemoglobina", "ferritina", "ist", "calcio", "fosforo", "pth",
    "potassio", "bicarbonato", "albumina", "hbsag", "anti_hbs", "anti_hcv"
]


@dataclass
class ExamItem:
    nome_padronizado: str
    nome_original: str
    valor: str
    unidade: str
    referencia: str
    data_coleta: str
    incerto: bool
    trecho_fonte: str


@dataclass
class MedicationItem:
    nome: str
    dose: str
    via: str
    frequencia: str
    fonte: str


@dataclass
class ClinicalCase:
    paciente: Dict[str, Any]
    contexto: Dict[str, Any]
    exames: List[Dict[str, Any]]
    medicacoes_em_uso: List[Dict[str, Any]]
    prescricao_dialitica: Dict[str, Any]
    evolucao_previa: Dict[str, Any]
    intercorrencias: List[Dict[str, Any]]
    problemas_ativos: List[Dict[str, Any]]
    pendencias: List[Dict[str, Any]]
    rastreabilidade: Dict[str, Any]


def normalize_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "á": "a", "à": "a", "â": "a", "ã": "a",
        "é": "e", "ê": "e",
        "í": "i",
        "ó": "o", "ô": "o", "õ": "o",
        "ú": "u",
        "ç": "c",
    }
    out = text.lower()
    for old, new in replacements.items():
        out = out.replace(old, new)
    return out


def standardize_exam_name(raw_name: str) -> str:
    name_norm = normalize_text(raw_name)
    for standard, synonyms in EXAM_SYNONYMS.items():
        for syn in synonyms:
            # Use exact match for short synonyms (e.g. "p", "k", "na") to avoid
            # false substring hits like "p" matching inside "pth".
            if name_norm == syn or (len(syn) > 2 and syn in name_norm):
                return standard
    return name_norm.replace(" ", "_")


def parse_date(value: str) -> str:
    if not value:
        return ""
    for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d"]:
        try:
            return datetime.strptime(value.strip(), fmt).date().isoformat()
        except Exception:
            continue
    return value.strip()


def basic_extract_exams(text: str, default_date: str = "") -> List[ExamItem]:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    exams: List[ExamItem] = []
    pattern = re.compile(
        r"^(?P<name>[A-Za-zÀ-ÿ0-9\-\s\/%]+?)\s*(?:\:|\=)?\s*"
        r"(?P<value>[<>]?\d+[\.,]?\d*)\s*"
        r"(?P<unit>[A-Za-zµ/%²0-9\-\^\.]+)?\s*$"
    )

    for line in lines:
        m = pattern.match(line)
        if not m:
            continue
        raw_name = m.group("name").strip(" -")
        value = m.group("value") or ""
        unit = (m.group("unit") or "").strip()
        std_name = standardize_exam_name(raw_name)

        exams.append(
            ExamItem(
                nome_padronizado=std_name,
                nome_original=raw_name,
                valor=value.replace(",", "."),
                unidade=unit,
                referencia="",
                data_coleta=default_date,
                incerto=False,
                trecho_fonte=line,
            )
        )
    return exams


def parse_medications(text: str) -> List[MedicationItem]:
    meds: List[MedicationItem] = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for line in lines:
        meds.append(MedicationItem(
            nome=line,
            dose="",
            via="",
            frequencia="",
            fonte="texto_manual"
        ))
    return meds


def exams_to_index(exams: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    indexed: Dict[str, Dict[str, Any]] = {}
    for exam in exams:
        key = exam.get("nome_padronizado", "")
        if key:
            indexed[key] = exam
    return indexed


def try_float(value: str) -> Optional[float]:
    try:
        return float(str(value).replace(",", "."))
    except Exception:
        return None


def reconciler(current_exams: List[Dict[str, Any]], previous_exams: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    current = exams_to_index(current_exams)
    previous = exams_to_index(previous_exams)
    findings: List[Dict[str, Any]] = []

    tracked = sorted(set(current.keys()) | set(previous.keys()))
    for key in tracked:
        curr = current.get(key)
        prev = previous.get(key)
        if curr and prev:
            curr_val = try_float(curr.get("valor", ""))
            prev_val = try_float(prev.get("valor", ""))
            trend = "sem tendencia definida"
            if curr_val is not None and prev_val is not None:
                if curr_val > prev_val:
                    trend = "aumentou"
                elif curr_val < prev_val:
                    trend = "reduziu"
                else:
                    trend = "estavel"
            findings.append({
                "dominio": key,
                "descricao": f"{key}: atual={curr.get('valor','')} {curr.get('unidade','')} | previo={prev.get('valor','')} {prev.get('unidade','')} | tendencia={trend}",
                "status": "comparado"
            })
        elif curr and not prev:
            findings.append({
                "dominio": key,
                "descricao": f"{key}: presente no material atual sem comparativo previo estruturado.",
                "status": "novo_no_contexto"
            })
        elif prev and not curr:
            findings.append({
                "dominio": key,
                "descricao": f"{key}: presente apenas no material previo; nao localizado no material atual.",
                "status": "ausente_no_atual"
            })
    return findings


def guideline_engine(case_data: ClinicalCase) -> List[Dict[str, Any]]:
    exams = exams_to_index(case_data.exames)
    suggestions: List[Dict[str, Any]] = []

    hb = try_float(exams.get("hemoglobina", {}).get("valor", ""))
    ferritin = try_float(exams.get("ferritina", {}).get("valor", ""))
    tsat = try_float(exams.get("ist", {}).get("valor", ""))
    phosphorus = try_float(exams.get("fosforo", {}).get("valor", ""))
    pth = try_float(exams.get("pth", {}).get("valor", ""))
    potassium = try_float(exams.get("potassio", {}).get("valor", ""))

    if hb is not None:
        lacunas = []
        if ferritin is None:
            lacunas.append("ferritina ausente")
        if tsat is None:
            lacunas.append("IST/TSAT ausente")
        suggestions.append({
            "problema": "anemia",
            "achado": f"Hemoglobina atual: {hb}",
            "base_documental": "Base guideline a ser conectada via RAG; placeholder para KDIGO anemia e protocolo local.",
            "lacunas": lacunas,
            "sugestao_prudente": "Revisar estrategia de anemia com correlacao longitudinal e protocolo local. Nao fechar prescricao automaticamente."
        })

    if phosphorus is not None or pth is not None:
        lacunas = []
        if exams.get("calcio") is None:
            lacunas.append("calcio ausente")
        if exams.get("fosfatase_alcalina") is None:
            lacunas.append("fosfatase alcalina ausente")
        suggestions.append({
            "problema": "disturbio_mineral_osseo",
            "achado": f"Fosforo={phosphorus if phosphorus is not None else 'NA'} | PTH={pth if pth is not None else 'NA'}",
            "base_documental": "Base guideline a ser conectada via RAG; placeholder para CKD-MBD e protocolo local.",
            "lacunas": lacunas,
            "sugestao_prudente": "Reavaliar tendencia, adesao e esquema atual antes de qualquer ajuste terapeutico."
        })

    if potassium is not None:
        severity = "valor critico" if potassium >= 6.0 else "sem criticidade automatica"
        suggestions.append({
            "problema": "potassio",
            "achado": f"Potassio atual: {potassium} ({severity})",
            "base_documental": "Placeholder de seguranca; exige revisao medica imediata se criticidade clinica.",
            "lacunas": [],
            "sugestao_prudente": "Checar contexto clinico, hemolise, ECG e necessidade de acao imediata conforme protocolo local."
        })

    return suggestions


def build_short_evolution(case_data: ClinicalCase, guideline_notes: List[Dict[str, Any]]) -> str:
    modality = case_data.contexto.get("modalidade_trs", "TRS nao informada")
    summary_parts = []
    problems = [g["problema"] for g in guideline_notes]

    if "anemia" in problems:
        summary_parts.append("achados relacionados a anemia em seguimento")
    if "disturbio_mineral_osseo" in problems:
        summary_parts.append("alteracoes de metabolismo mineral/osseo")
    if "potassio" in problems:
        summary_parts.append("necessidade de correlacao clinica do potassio")

    if not summary_parts:
        summary_parts.append("sem achados automaticamente categorizados pelo modelo inicial")

    return (
        f"Paciente em {modality}. Material atual mostra {', '.join(summary_parts)}. "
        f"Minuta automatizada para revisao medica, sem fechamento autonomo de prescricao."
    )


def build_detailed_evolution(case_data: ClinicalCase, guideline_notes: List[Dict[str, Any]], reconciliation: List[Dict[str, Any]]) -> str:
    lines = []
    modality = case_data.contexto.get("modalidade_trs", "nao informada")
    lines.append(f"Paciente em {modality}.")

    if guideline_notes:
        lines.append("Achados estruturados:")
        for item in guideline_notes:
            lines.append(f"- {item['problema']}: {item['achado']}")

    if reconciliation:
        lines.append("Comparacao longitudinal:")
        for item in reconciliation[:8]:
            lines.append(f"- {item['descricao']}")

    pend = case_data.pendencias
    if pend:
        lines.append("Pendencias:")
        for item in pend:
            lines.append(f"- {item.get('descricao','')}")

    lines.append("Texto gerado para revisao humana final.")
    return "\n".join(lines)


def safety_auditor(case_data: ClinicalCase, detailed_note: str) -> Dict[str, Any]:
    exams = exams_to_index(case_data.exames)
    issues: List[str] = []

    if not case_data.contexto.get("modalidade_trs"):
        issues.append("Modalidade de TRS ausente.")

    for name in ["hemoglobina", "potassio", "fosforo", "calcio", "pth"]:
        exam = exams.get(name)
        if exam and not exam.get("unidade"):
            issues.append(f"Exame critico sem unidade: {name}.")

    potassium = try_float(exams.get("potassio", {}).get("valor", ""))
    if potassium is not None and potassium >= 6.0 and "potassio" not in normalize_text(detailed_note):
        issues.append("Possivel hiperpotassemia sem destaque adequado na narrativa.")

    if issues:
        status = "Bloqueado" if len(issues) >= 2 else "Aprovado com ressalvas"
    else:
        status = "Aprovado"

    return {
        "status": status,
        "ressalvas": issues,
    }


def default_case() -> ClinicalCase:
    return ClinicalCase(
        paciente={"nome": "", "registro": "", "data_nascimento": "", "sexo": ""},
        contexto={
            "modalidade_trs": "HD",
            "local_atendimento": "",
            "data_referencia": datetime.now().date().isoformat(),
            "fonte_entrada": [],
        },
        exames=[],
        medicacoes_em_uso=[],
        prescricao_dialitica={
            "modalidade": "",
            "frequencia": "",
            "duracao_sessao": "",
            "banho": "",
            "anticoagulacao": "",
            "uf_programada": "",
            "peso_seco": "",
            "acesso": "",
        },
        evolucao_previa={"data": "", "texto": ""},
        intercorrencias=[],
        problemas_ativos=[],
        pendencias=[],
        rastreabilidade={
            "documentos_entrada": [],
            "guidelines_consultadas": [],
            "versao_modelo": "mvp-streamlit-v0.1",
            "data_processamento": datetime.now().isoformat(),
        },
    )


def add_basic_pendencies(case_data: ClinicalCase) -> None:
    exams = exams_to_index(case_data.exames)
    pending = []
    for key in IMPORTANT_EXAMS:
        if key not in exams:
            pending.append({"tipo": "exame_ausente", "descricao": f"Exame nao localizado: {key}"})
    case_data.pendencias = pending


def export_docx(short_note: str, detailed_note: str, audit: Dict[str, Any]) -> Optional[bytes]:
    if docx is None:
        return None

    document = docx.Document()
    document.add_heading("TRS Scribe Assist - Relatorio", level=1)
    document.add_paragraph(f"Status da auditoria: {audit.get('status', '')}")

    document.add_heading("Evolucao curta", level=2)
    document.add_paragraph(short_note)

    document.add_heading("Evolucao detalhada", level=2)
    document.add_paragraph(detailed_note)

    if audit.get("ressalvas"):
        document.add_heading("Ressalvas", level=2)
        for item in audit["ressalvas"]:
            document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
